"""
Generator Dokumen Admin Guru MI (KBC & KMA 1503/2025)
--------------------------------------------------------------------------------
Pembaruan V4.13 (Joyful Edition): 
Perubahan total pada UI Docx Color Palette menjadi lebih ceria, warna-warni (vibrant & pastel), 
mendukung prinsip Joyful Learning secara visual tanpa mengurangi formalitas dokumen.
"""

import io
import json
import re
import datetime
import time

import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ==============================================================================
# PALET WARNA JOYFUL (VIBRANT & PASTEL)
# ==============================================================================
COLOR_TITLE = "6C5CE7"       # Vibrant Purple (Ungu Ceria)
COLOR_IDENTITY_HEAD = "0984E3"   # Bright Blue (Biru Terang)
COLOR_LABEL = "DEEAF1"       # Soft Blue Label
COLOR_VALUE = "FFFFFF"       # Putih
COLOR_SECTION_A = "00B894"   # Mint Green (Hijau Segar)
COLOR_SECTION_B = "E17055"   # Warm Coral (Oranye Hangat)
COLOR_MEETING = "F39C12"     # Sunny Orange
COLOR_LAMPIRAN_I = "D63031"  # Vibrant Red
COLOR_LAMPIRAN_II = "00CEC9" # Bright Teal (Tosca)
COLOR_LAMPIRAN_III = "E84393" # Playful Pink (Merah Muda Ceria)
COLOR_LAMPIRAN_V = "F39C12"  # Sunny Yellow/Orange

MODEL_NAME = "google/diffusiongemma-26b-a4b-it"
NVIDIA_BASE_URL = "https://integrate.api.nvidia.com/v1"

JENJANG_FASE = {
    "RA/TK (Fase Fondasi)": "Fondasi", "Kelas 1 SD/MI (Fase A)": "A", "Kelas 2 SD/MI (Fase A)": "A",
    "Kelas 3 SD/MI (Fase B)": "B", "Kelas 4 SD/MI (Fase B)": "B", "Kelas 5 SD/MI (Fase C)": "C",
    "Kelas 6 SD/MI (Fase C)": "C", "Kelas 7 SMP/MTs (Fase D)": "D", "Kelas 8 SMP/MTs (Fase D)": "D",
    "Kelas 9 SMP/MTs (Fase D)": "D", "Kelas 10 SMA/MA (Fase E)": "E", "Kelas 11 SMA/MA (Fase F)": "F",
    "Kelas 12 SMA/MA (Fase F)": "F",
}

st.set_page_config(page_title="MIFSAL ADMIN GURU V4.13", page_icon="🎨", layout="wide")

@st.cache_resource
def get_client():
    api_key = st.secrets.get("NVIDIA_API_KEY")
    if not api_key:
        st.error("NVIDIA_API_KEY belum ada di st.secrets.")
        st.stop()
    return OpenAI(base_url=NVIDIA_BASE_URL, api_key=api_key)

def call_ai(prompt: str, temperature=0.7) -> dict:
    client = get_client()
    text = ""
    max_retries = 4
    
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model=MODEL_NAME, messages=[{"role": "user", "content": prompt}],
                temperature=temperature, max_tokens=16192,
            )
            raw_content = response.choices[0].message.content
            if raw_content is None:
                if attempt < max_retries - 1:
                    time.sleep(5)
                    continue
                else: text = "" 
            else:
                text = raw_content.strip()
                break
        except Exception as e:
            if "429" in str(e) or "Too Many Requests" in str(e):
                if attempt < max_retries - 1:
                    time.sleep(15)
                    continue
            raise e

    st.session_state["raw_ai_output"] = text 
    text = text.replace("```json", "").replace("```", "").strip()
    start = text.find('{')
    end = text.rfind('}')
    if start != -1 and end != -1: text = text[start:end+1]
        
    text = text.replace('\n', ' ').replace('\r', '')
    text = re.sub(r'[\x00-\x1f]', '', text)
    text = re.sub(r',\s*([}\]])', r'\1', text)
    
    try: return json.loads(text)
    except json.JSONDecodeError as e: return {}

def get_aggregated_sinkronisasi_context(all_chapters_data):
    sinkron_text = ""
    for idx, chap in enumerate(all_chapters_data):
        bab_name = chap["bab"]
        d1 = chap["d1"]; d2 = chap["d2"]
        sinkron_text += f"\n--- INFORMASI BAB {idx+1}: {bab_name} ---\n"
        if d1 and "desain" in d1:
            cp = d1["desain"].get("capaian_pembelajaran", "")
            tp = d1["desain"].get("tujuan_pembelajaran", [])
            if isinstance(tp, list): tp = ", ".join(tp)
            sinkron_text += f"CP: '{cp}'\nTP: '{tp}'\n"
        if d2 and "pertemuan" in d2:
            materi_list = [f"Pertemuan {p.get('nomor', '')}: {p.get('materi', '')}" for p in d2.get("pertemuan", []) if isinstance(p, dict)]
            materi_str = " | ".join(materi_list)
            sinkron_text += f"Rincian Materi & Sub-Bab: {materi_str}\n"

    if sinkron_text:
        return f"\n\n[SANGAT PENTING - MASTER SINKRONISASI DOKUMEN]\nKamu WAJIB mematuhi data dari seluruh Bab berikut agar semua dokumen selaras:\n{sinkron_text}\nRangkum dan susun secara logis sesuai Sub-Bab yang tercantum."
    return ""

# ==============================================================================
# PROMPT MODUL AJAR 
# ==============================================================================
def prompt_step_1(form, bab_name):
    return f"""Kamu pakar Kurikulum Merdeka Pendekatan Deep Learning Berbasis Cinta (KBC). 
Mapel: {form['mapel']}, Jenjang: {form['kelas']}, Topik Utama dan Sub-Bab: {bab_name}. 
PENTING: CP dan TP WAJIB mengacu "KMA 1503 Tahun 2025". Masukkan rincian Sub-Bab ke dalam 'topik_pembelajaran'. Isi Dimensi Profil Kelulusan dengan 4 poin bernomor. Isi Topik Panca Cinta secara spesifik.
Balas HANYA JSON:
{{"identifikasi": {{"pengetahuan_awal": ["str"], "minat_belajar": ["str"], "latar_belakang": "str (1 paragraf panjang)", "kebutuhan_belajar": ["str"], "dimensi_profil": ["str"], "panca_cinta": ["str"]}}, "desain": {{"capaian_pembelajaran": "str", "tujuan_pembelajaran": ["str"], "lintas_disiplin": ["str"], "topik_pembelajaran": ["Rincian Sub-Bab dari judul..."], "praktik_pedagogi": ["str"], "lingkungan_belajar": ["str"], "kemitraan_pembelajaran": ["str"], "pemanfaatan_digital": ["str"]}}}}"""

def prompt_step_2(form, bab_name, jumlah_pertemuan, step1):
    return f"""Lanjutkan modul {form['mapel']} bab {bab_name}. Buat Pengalaman Belajar untuk TEPAT {jumlah_pertemuan} pertemuan. 
[ATURAN MUTLAK]: Distribusikan/bagi materi berdasarkan rincian Sub-Bab yang ada pada judul secara berurutan ke dalam {jumlah_pertemuan} pertemuan tersebut. 'materi' setiap pertemuan HARUS mencantumkan nama Sub-Bab yang sedang dibahas.
Setiap pertemuan HARUS memiliki:
1. Fase Pembukaan (Waktu, Aktivitas bernomor)
2. Fase Inti: Dibagi 3 tahap (Memahami, Mengaplikasikan, Merefleksikan). Tiap tahap memiliki 'sintak_pbl' dan aktivitas bernomor.
3. Fase Penutup (Waktu, Aktivitas bernomor).
Balas HANYA JSON:
{{"pertemuan": [{{"nomor": 1, "materi": "Sub-Bab Spesifik Pertemuan Ini", "durasi": "1 x 35 Menit", "pembukaan": {{"waktu": "5'", "aktivitas": ["str", "str"]}}, "inti_memahami": {{"sintak_pbl": "Langkah 1: Orientasi Masalah", "aktivitas": ["str", "str"]}}, "inti_mengaplikasikan": {{"sintak_pbl": "Langkah 2: Organisasi Belajar", "aktivitas": ["str", "str"]}}, "inti_merefleksikan": {{"sintak_pbl": "Langkah 5: Analisis & Evaluasi", "aktivitas": ["str", "str"]}}, "penutup": {{"waktu": "5'", "aktivitas": ["str", "str"]}}}}]}}"""

def prompt_step_3(form, bab_name, jumlah_pertemuan, step2):
    return f"""Tahap akhir modul {form['mapel']} bab {bab_name}. 
Pastikan bagian 'materi_ajar' (Sesi 1, 2, dst) dan 'lkpd' membahas spesifik Sub-Bab yang telah didistribusikan.
Sediakan Asesmen, Rubrik Pengetahuan (4 Aspek dengan skor 4,3,2,1), dan 5 Soal PG Remedial.
Balas HANYA JSON:
{{"penilaian": {{"awal": ["str"], "formatif_as": ["str"], "formatif_for": ["str"], "sumatif": ["str"]}}, "asesmen_awal_lisan": ["Soal 1", "Soal 2"], "rubrik_pengetahuan": [{{"aspek": "Aspek 1", "skor_4": "str", "skor_3": "str", "skor_2": "str", "skor_1": "str"}}], "sumatif_hots": ["Soal HOTS 1", "Soal HOTS 2", "Soal HOTS 3", "Soal HOTS 4", "Soal HOTS 5"], "materi_ajar": [{{"sesi": "SESI 1: Nama Sub-Bab", "tujuan": "str", "pengantar": "str", "konsep_utama": "str", "contoh": "str"}}], "lkpd": [{{"pertemuan": 1, "judul": "Judul LKPD Sub-Bab Tersebut", "memahami": ["Soal 1", "Soal 2"], "mengaplikasikan": "Tugas aplikasi", "merefleksikan": "Pertanyaan refleksi"}}], "remedial_pg": [{{"soal": "Pertanyaan PG?", "a": "Opsi A", "b": "Opsi B", "c": "Opsi C", "d": "Opsi D", "kunci": "a"}}], "pengayaan": ["Tugas 1", "Tugas 2"], "glosarium": [{{"istilah": "str", "definisi": "str"}}], "daftar_pustaka": ["Referensi 1", "Referensi 2"]}}"""


# ==============================================================================
# PROMPT KHUSUS LKPD (PERTEMUAN)
# ==============================================================================
def prompt_lkpd_interaktif_per_pertemuan(form, bab_name, pertemuan_ke):
    return f"""Buat LKPD Interaktif untuk Mapel {form['mapel']} {form['kelas']} Topik {bab_name}, KHUSUS UNTUK PERTEMUAN KE-{pertemuan_ke} SAJA.
Pastikan materi LKPD ini nyambung dengan urutan Sub-Bab untuk pertemuan tersebut.
[ATURAN EMAS LKPD]: LKPD HARUS INTERAKTIF DENGAN SKENARIO DUNIA NYATA.
Balas HANYA JSON:
{{"lkpd": {{"pertemuan": {pertemuan_ke}, "topik_judul": "Topik Pertemuan {pertemuan_ke}: [Nama Sub-Bab]", "tujuan_pembelajaran": ["Memahami..."], "aktivitas": [{{"judul_aktivitas": "Aktivitas 1: Kisah/Konteks", "konteks": "Bayangkan kamu sedang...", "ayo_berpikir": ["Langkah 1...", "Langkah 2..."], "koneksi_matematika": ["Kalimat matematika...", "Bandingkan..."], "kesimpulan": "Artinya adalah...", "latihan_berpikir": "Kasus untuk siswa..."}}]}}}}"""


# ==============================================================================
# PROMPT DOKUMEN LAIN (GLOBAL & PER BAB)
# ==============================================================================
def prompt_cp(form, aggregated_context):
    return f"""Buat Capaian Pembelajaran (CP) Mapel {form['mapel']} Fase/Kelas {form['kelas']}. KONTEN CP WAJIB MENGACU KMA 1503 TAHUN 2025. {aggregated_context}
Balas HANYA JSON: {{"rasional": "str 1 paragraf", "tujuan": ["str"], "elemen": [{{"nama": "str", "deskripsi": "str"}}], "cp_paragraf": "str 1 paragraf", "cp_tabel": [{{"elemen": "str", "capaian": "str"}}]}}"""

def prompt_atp(form, aggregated_context):
    return f"""Buat Alur Tujuan Pembelajaran (ATP) Mapel {form['mapel']} {form['kelas']}. WAJIB MENGACU KMA 1503 TAHUN 2025. {aggregated_context}
Balas HANYA JSON: {{"cp_fase": "str", "rows": [{{"no": "1", "elemen": "str", "tp": "str", "atp": "str", "materi": "str", "jp": "str"}}]}}"""

def prompt_prota(form, aggregated_context):
    return f"""Buat Program Tahunan (PROTA) Mapel {form['mapel']} {form['kelas']}. Tuliskan nama Bab beserta Sub-babnya dari referensi berikut. {aggregated_context}
Balas HANYA JSON: {{"rows": [{{"semester": "1", "no": "1", "materi": "Bab 1... (Sub-Bab...)", "jp": "str", "keterangan": "str"}}]}}"""

def prompt_promes(form, aggregated_context):
    sem = form['semester']
    if "1 & 2" in sem:
        bulan = "Juli, Agustus, September, Oktober, November, Desember, Januari, Februari, Maret, April, Mei, Juni"
    elif "1" in sem:
        bulan = "Juli, Agustus, September, Oktober, November, Desember"
    else:
        bulan = "Januari, Februari, Maret, April, Mei, Juni"
        
    return f"""Buat Program Semester Lengkap Mapel {form['mapel']} {form['kelas']}. 
Semester yang dipilih: {sem}.
[ATURAN PENTING]: Ekstrak rincian Sub-Bab dari master konteks dan jadikan baris dengan 'jenis': 'sub'.
{aggregated_context}
Balas HANYA JSON dengan struktur INI:
{{
  "tables": [
    {{
      "judul_semester": "SEMESTER 1 (GANJIL)",
      "bulan": ["Juli", "Agustus", "September", "Oktober", "November", "Desember"],
      "rows": [
         {{"jenis": "bab", "materi": "Bab 1: Nama Bab", "jp": "30", "distribusi": []}},
         {{"jenis": "sub", "materi": "A. Nama Sub-Bab 1", "jp": "10", "distribusi": [{{"bulan": "Juli", "minggu": 1, "jp": "4"}}]}},
         {{"jenis": "sub", "materi": "B. Nama Sub-Bab 2", "jp": "10", "distribusi": [{{"bulan": "Juli", "minggu": 2, "jp": "4"}}]}}
      ]
    }}
  ]
}}"""

def prompt_kktp_per_bab(form, bab_name, tp_list):
    return f"""Buat Kriteria Ketercapaian Tujuan Pembelajaran (KKTP) Mapel {form['mapel']} {form['kelas']}.
KHUSUS UNTUK BAB INI: {bab_name}. Tujuan Pembelajaran: {tp_list}
Buatlah tabel indikator KKTP yang sangat lengkap.
Balas HANYA JSON: {{"rows": [{{"tp": "str", "kriteria": "str"}}]}}"""


def generate_jurnal_otomatis(form, all_chapters_data):
    jurnal_rows = []
    pertemuan_global = 1
    for chap in all_chapters_data:
        bab_name = chap.get("bab", "Materi")
        d2 = chap.get("d2", {})
        pertemuan_list = safe_list(d2.get("pertemuan", []))
        for p in pertemuan_list:
            if not isinstance(p, dict): continue
            topik_materi = p.get("materi", bab_name) # Sudah memuat Sub-Bab
            akt = p.get("inti_mengaplikasikan", {}).get("aktivitas", ["Pembelajaran KBC"])
            aktivitas_utama = str(akt[0]) if isinstance(akt, list) and len(akt) > 0 else str(akt)
            jurnal_rows.append({
                "pertemuan": str(pertemuan_global),
                "topik": topik_materi,
                "aktivitas": aktivitas_utama,
                "asesmen": "Formatif/Lisan/Penugasan LKPD"
            })
            pertemuan_global += 1
    return {"rows": jurnal_rows}


# ==============================================================================
# FUNGSI PEMBANTU FORMATTING DOCX
# ==============================================================================
def safe_list(val, default=None):
    if default is None: default = ["-"]
    if val is None: return default
    if isinstance(val, str): return [val]
    if isinstance(val, list) and len(val) > 0: return val
    return default

def set_cell_background(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def style_cell(cell, text, bold=False, color="000000", center=False, size=10, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = p.add_run(str(text))
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)

def banner(doc, text, hex_color, size=12):
    table = doc.add_table(rows=1, cols=1); table.autofit = True; table.columns[0].width = Cm(18)
    cell = table.rows[0].cells[0]
    set_cell_background(cell, hex_color)
    style_cell(cell, text, bold=True, color="FFFFFF", size=size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def field_table(doc):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4.5); table.columns[1].width = Cm(13.5)
    return table

def add_field_row(table, label, content_items):
    row = table.add_row()
    label_cell, value_cell = row.cells[0], row.cells[1]
    label_cell.width = Cm(4.5); value_cell.width = Cm(13.5)
    set_cell_background(label_cell, COLOR_LABEL); set_cell_background(value_cell, COLOR_VALUE)
    style_cell(label_cell, label, bold=True, size=10.5)
    value_cell.text = ""
    if isinstance(content_items, (list, tuple)):
        for i, item in enumerate(content_items):
            p = value_cell.paragraphs[0] if i == 0 else value_cell.add_paragraph()
            p.style = value_cell.paragraphs[0].style
            p.add_run(f"\u2022 {item}").font.size = Pt(10.5)
    else:
        value_cell.paragraphs[0].add_run(str(content_items)).font.size = Pt(10.5)
    return row

def create_base_doc(landscape=False):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.5); section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)
    if landscape: section.orientation = 1; section.page_width, section.page_height = section.page_height, section.page_width
    return doc

def add_signatures(doc, form, full_width=False):
    doc.add_paragraph("\n")
    sig_table = doc.add_table(rows=1, cols=2)
    w = Cm(12) if full_width else Cm(9)
    sig_table.columns[0].width = w; sig_table.columns[1].width = w
    p1 = sig_table.rows[0].cells[0].paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run(f"Mengetahui,\nKepala Sekolah {form['sekolah']}\n\n\n\n")
    p1.add_run(f"({form['kepala_madrasah']})").bold = True
    p1.add_run("\nNIP. .....................................")
    
    p2 = sig_table.rows[0].cells[1].paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(f"{form['titimangsa']}\nGuru Mata Pelajaran\n\n\n\n")
    p2.add_run(f"({form['penyusun']})").bold = True
    p2.add_run("\nNIP. .....................................")

def create_header(doc, title, form):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title); run.bold = True; run.font.size = Pt(14)
    doc.add_paragraph(f"Mata Pelajaran: {form['mapel']}")
    doc.add_paragraph(f"Nama Sekolah: {form['sekolah']}")
    if "Penyusun" in title or "ATP" in title or "KKTP" in title: doc.add_paragraph(f"Nama Penyusun: {form['penyusun']}")
    doc.add_paragraph(f"Fase/Kelas: {form['kelas']}")
    doc.add_paragraph(f"Tahun Ajaran: {form['tahun_pelajaran']}")
    doc.add_paragraph()


# ==============================================================================
# BUILDERS (Global)
# ==============================================================================
def build_cp(form, ai_data):
    doc = create_base_doc(landscape=False)
    create_header(doc, "CAPAIAN PEMBELAJARAN (CP)", form)
    
    doc.add_heading("A. Rasional Mata Pelajaran", level=3)
    doc.add_paragraph(ai_data.get("rasional", ""))
    
    doc.add_heading("B. Tujuan Mata Pelajaran", level=3)
    for t in safe_list(ai_data.get("tujuan", [])): 
        doc.add_paragraph(t, style='List Bullet')
        
    doc.add_heading("C. Elemen-elemen Mata Pelajaran", level=3)
    t_elemen = doc.add_table(rows=1, cols=2)
    t_elemen.style = 'Table Grid'
    set_cell_background(t_elemen.cell(0, 0), "D6EAF8"); style_cell(t_elemen.cell(0, 0), "Elemen", bold=True)
    set_cell_background(t_elemen.cell(0, 1), "D6EAF8"); style_cell(t_elemen.cell(0, 1), "Deskripsi", bold=True)
    t_elemen.columns[0].width = Cm(4.0); t_elemen.columns[1].width = Cm(14.0)
    for el in safe_list(ai_data.get("elemen", [])):
        if isinstance(el, dict):
            r = t_elemen.add_row().cells
            style_cell(r[0], el.get("nama", "")); style_cell(r[1], el.get("deskripsi", ""))
    doc.add_paragraph()
    
    doc.add_heading(f"D. Capaian Pembelajaran Fase", level=3)
    doc.add_paragraph(ai_data.get("cp_paragraf", ""))
    t_cp = doc.add_table(rows=1, cols=2)
    t_cp.style = 'Table Grid'
    set_cell_background(t_cp.cell(0, 0), "D6EAF8"); style_cell(t_cp.cell(0, 0), "Elemen", bold=True)
    set_cell_background(t_cp.cell(0, 1), "D6EAF8"); style_cell(t_cp.cell(0, 1), "Capaian Pembelajaran", bold=True)
    t_cp.columns[0].width = Cm(4.0); t_cp.columns[1].width = Cm(14.0)
    for cp in safe_list(ai_data.get("cp_tabel", [])):
        if isinstance(cp, dict):
            r = t_cp.add_row().cells
            style_cell(r[0], cp.get("elemen", "")); style_cell(r[1], cp.get("capaian", ""))
            
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()

def build_atp(form, ai_data):
    doc = create_base_doc(landscape=True)
    create_header(doc, "ALUR TUJUAN PEMBELAJARAN (ATP)", form)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    headers = ["No.", "Elemen", "Tujuan Pembelajaran (TP) per Bab", "Alur Tujuan Pembelajaran (ATP)", "Materi Pokok", "Alokasi Waktu (JP)"]
    for i, h in enumerate(headers):
        set_cell_background(table.cell(0, i), "FCF3CF"); style_cell(table.cell(0, i), h, bold=True, center=True)
    table.columns[0].width = Cm(1.0); table.columns[5].width = Cm(2.0)
    
    for row in safe_list(ai_data.get("rows"), []):
        if not isinstance(row, dict): continue
        r = table.add_row().cells
        style_cell(r[0], row.get("no", ""), center=True); style_cell(r[1], row.get("elemen", ""))
        style_cell(r[2], row.get("tp", "")); style_cell(r[3], row.get("atp", ""))
        style_cell(r[4], row.get("materi", "")); style_cell(r[5], row.get("jp", ""), center=True)
        
    add_signatures(doc, form, full_width=True)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()

def build_prota(form, ai_data):
    doc = create_base_doc(landscape=False)
    create_header(doc, "PROGRAM TAHUNAN (PROTA)", form)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    headers = ["Semester", "No", "Materi Pokok / Bab", "Alokasi Waktu (JP)", "Keterangan"]
    for i, h in enumerate(headers):
        set_cell_background(table.cell(0, i), "D1F2EB"); style_cell(table.cell(0, i), h, bold=True, center=True)
    table.columns[0].width = Cm(2.5); table.columns[1].width = Cm(1.0); table.columns[3].width = Cm(3.0)
    
    for row in safe_list(ai_data.get("rows"), []):
        if not isinstance(row, dict): continue
        r = table.add_row().cells
        style_cell(r[0], row.get("semester", ""), center=True); style_cell(r[1], row.get("no", ""), center=True)
        style_cell(r[2], row.get("materi", "")); style_cell(r[3], row.get("jp", ""), center=True)
        style_cell(r[4], row.get("keterangan", ""), center=True)
        
    add_signatures(doc, form)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()

def build_promes(form, ai_data):
    doc = create_base_doc(landscape=True)
    create_header(doc, "PROGRAM SEMESTER (PROSEM)", form)
    
    tables_data = safe_list(ai_data.get("tables", []))
    for t_data in tables_data:
        if not isinstance(t_data, dict): continue
        judul_sem = t_data.get("judul_semester", "SEMESTER")
        doc.add_heading(judul_sem, level=2)
        bulan_list = safe_list(t_data.get("bulan", []))
        if not bulan_list: continue
        
        total_cols = 2 + (len(bulan_list) * 5)
        table = doc.add_table(rows=2, cols=total_cols)
        table.style = 'Table Grid'
        
        table.cell(0, 0).merge(table.cell(1, 0)); style_cell(table.cell(0, 0), "Materi / Tujuan Pembelajaran", bold=True, center=True)
        table.cell(0, 1).merge(table.cell(1, 1)); style_cell(table.cell(0, 1), "JP", bold=True, center=True)
        set_cell_background(table.cell(0, 0), "FADBD8"); set_cell_background(table.cell(0, 1), "FADBD8")
        table.columns[0].width = Cm(6.0); table.columns[1].width = Cm(1.5)
        
        col_idx = 2
        for b in bulan_list:
            table.cell(0, col_idx).merge(table.cell(0, col_idx + 4))
            style_cell(table.cell(0, col_idx), b, bold=True, center=True)
            set_cell_background(table.cell(0, col_idx), "FADBD8")
            for w in range(5):
                style_cell(table.cell(1, col_idx + w), str(w + 1), bold=True, center=True)
                set_cell_background(table.cell(1, col_idx + w), "FDEDEC")
                table.columns[col_idx + w].width = Cm(0.6)
            col_idx += 5
            
        for row in safe_list(t_data.get("rows", [])):
            if not isinstance(row, dict): continue
            r = table.add_row().cells
            jenis = row.get("jenis", "sub"); materi = row.get("materi", ""); jp = str(row.get("jp", ""))
            is_bab = (jenis == "bab")
            
            style_cell(r[0], materi, bold=is_bab); style_cell(r[1], jp, bold=is_bab, center=True)
            if is_bab:
                set_cell_background(r[0], "FCF3CF"); set_cell_background(r[1], "FCF3CF")
                
            distribusi = safe_list(row.get("distribusi", []))
            for dist in distribusi:
                if not isinstance(dist, dict): continue
                d_bulan = dist.get("bulan", ""); d_minggu = int(dist.get("minggu", 1)); d_jp = str(dist.get("jp", ""))
                try:
                    b_idx = bulan_list.index(d_bulan)
                    c_idx = 2 + (b_idx * 5) + (d_minggu - 1)
                    if 2 <= c_idx < total_cols: style_cell(r[c_idx], d_jp, center=True)
                except ValueError: pass
        doc.add_paragraph()
        
    add_signatures(doc, form, full_width=True)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()

def build_kktp(form, combined_kktp_data):
    doc = create_base_doc(landscape=False)
    create_header(doc, "KRITERIA KETERCAPAIAN TUJUAN PEMBELAJARAN (KKTP)", form)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    for i, h in enumerate(["Tujuan Pembelajaran (TP)", "Kriteria Ketercapaian (Indikator)"]):
        set_cell_background(table.cell(0, i), "E8DAEF"); style_cell(table.cell(0, i), h, bold=True, center=True)
    
    for row in safe_list(combined_kktp_data.get("rows"), []):
        if not isinstance(row, dict): continue
        r = table.add_row().cells
        style_cell(r[0], row.get("tp", "")); style_cell(r[1], row.get("kriteria", ""))
    
    add_signatures(doc, form)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()

def build_jurnal(form, ai_data):
    doc = create_base_doc(landscape=False)
    create_header(doc, "JURNAL MENGAJAR HARIAN (GABUNGAN)", form)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    for i, h in enumerate(["Pertemuan", "Topik / Materi", "Aktivitas", "Asesmen", "Ket/Paraf"]):
        set_cell_background(table.cell(0, i), COLOR_TITLE); style_cell(table.cell(0, i), h, bold=True, color="FFFFFF", center=True)
        
    for row in safe_list(ai_data.get("rows"), []):
        if not isinstance(row, dict): continue
        r = table.add_row().cells
        style_cell(r[0], row.get("pertemuan", ""), center=True)
        style_cell(r[1], row.get("topik", "")); style_cell(r[2], row.get("aktivitas", ""))
        style_cell(r[3], row.get("asesmen", "")); style_cell(r[4], "") 
        
    add_signatures(doc, form)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


# ==============================================================================
# BUILDERS (Cover, Modul, dan LKPD Per Bab)
# ==============================================================================
def build_cover(form: dict, jenis_cover: str) -> bytes:
    is_landscape = (jenis_cover in ["Cover Program Tahunan & Semester", "Cover CP", "Cover ATP"])
    doc = create_base_doc(landscape=is_landscape)
    for _ in range(4): doc.add_paragraph()
    judul_utama = "BUKU PERANGKAT PEMBELAJARAN\n"
    if jenis_cover == "Cover Modul Ajar": judul_utama = "MODUL AJAR\nKURIKULUM BERBASIS CINTA\n"
    elif jenis_cover == "Cover Program Tahunan & Semester": judul_utama = "PROGRAM TAHUNAN DAN SEMESTER\n"
    elif jenis_cover == "Cover Jurnal Mengajar": judul_utama = "JURNAL MENGAJAR HARIAN\n"
    elif jenis_cover == "Cover CP": judul_utama = "CAPAIAN PEMBELAJARAN (CP)\n"
    elif jenis_cover == "Cover ATP": judul_utama = "ALUR TUJUAN PEMBELAJARAN (ATP)\n"
    elif jenis_cover == "Cover KKTP": judul_utama = "KRITERIA KETERCAPAIAN TUJUAN PEMBELAJARAN (KKTP)\n"
    
    p1 = doc.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(judul_utama); r1.bold = True; r1.font.size = Pt(22)
    r1_sub = p1.add_run("(Pendekatan Deep Learning - KMA 1503/2025)"); r1_sub.font.size = Pt(14)
    for _ in range(3): doc.add_paragraph()
    
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Mata Pelajaran : {form['mapel']}\n"); r2.bold = True; r2.font.size = Pt(16)
    r2_sub = p2.add_run(f"Kelas / Fase : {form['kelas']}\nSemester : {form['semester']}")
    r2_sub.font.size = Pt(14)
    for _ in range(5): doc.add_paragraph()
    
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3_a = p3.add_run("Disusun Oleh:\n"); r3_a.font.size = Pt(14)
    r3_b = p3.add_run(f"{form['penyusun']}"); r3_b.bold = True; r3_b.font.size = Pt(16)
    for _ in range(6): doc.add_paragraph()
    
    p4 = doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(f"{form['sekolah']}\n"); r4.bold = True; r4.font.size = Pt(18)
    r4_sub = p4.add_run(f"Tahun Pelajaran {form['tahun_pelajaran']}"); r4_sub.bold = True; r4_sub.font.size = Pt(14)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()

def build_modul_ajar(form: dict, bab_name: str, jumlah_pertemuan: int, full_data: dict) -> bytes:
    doc = create_base_doc(landscape=False)
    d1 = full_data.get("d1", {}); d2 = full_data.get("d2", {}); d3 = full_data.get("d3", {})

    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_head1 = p_head.add_run("MODUL AJAR\nKURIKULUM BERBASIS CINTA – PENDEKATAN DEEP LEARNING\n")
    r_head1.bold = True
    r_head1.font.size = Pt(14)
    r_head2 = p_head.add_run("Berdasarkan Capaian Pembelajaran Terbaru 2025 No.46\n")
    r_head2.font.size = Pt(11)

    doc.add_paragraph("IDENTITAS MODUL AJAR").bold = True
    t_id = doc.add_table(rows=5, cols=4); t_id.style = 'Table Grid'
    t_id.columns[0].width = Cm(4.0); t_id.columns[1].width = Cm(5.0)
    t_id.columns[2].width = Cm(4.0); t_id.columns[3].width = Cm(5.0)
    id_data = [
        ("Mata Pelajaran", form["mapel"], "Kelas / Fase", form["kelas"]),
        ("Semester", form["semester"], "Alokasi Waktu", f"{jumlah_pertemuan} Pertemuan × {form['alokasi']}"),
        ("Bab / Topik", bab_name, "", ""),
        ("Model Pembelajaran", "PBL (Problem Based Learning)", "Metode Pembelajaran", "Ceramah Interaktif, Diskusi Kelompok, Tanya Jawab"),
        ("Penyusun", form["penyusun"], "Sekolah", form["sekolah"]),
        ("Tahun Pelajaran", form["tahun_pelajaran"], "", "")
    ]
    for i, (k1, v1, k2, v2) in enumerate(id_data):
        if i >= len(t_id.rows): t_id.add_row()
        row = t_id.rows[i].cells
        style_cell(row[0], k1, bold=True); set_cell_background(row[0], "D6EAF8")
        style_cell(row[1], v1)
        if k2 == "": row[1].merge(row[3])
        else:
            style_cell(row[2], k2, bold=True); set_cell_background(row[2], "D6EAF8")
            style_cell(row[3], v2)
    doc.add_paragraph()

    doc.add_paragraph("A. IDENTIFIKASI PESERTA DIDIK").bold = True
    t_ident = doc.add_table(rows=0, cols=2); t_ident.style = 'Table Grid'
    t_ident.columns[0].width = Cm(5.0); t_ident.columns[1].width = Cm(13.0)
    ident_src = d1.get("identifikasi", {}) if isinstance(d1.get("identifikasi"), dict) else {}
    for label, key in [("Pengetahuan Awal", "pengetahuan_awal"), ("Minat Belajar", "minat_belajar"), ("Latar Belakang", "latar_belakang"), ("Kebutuhan Belajar", "kebutuhan_belajar"), ("Dimensi Profil Kelulusan (DPL)", "dimensi_profil"), ("Topik Panca Cinta", "panca_cinta")]:
        r = t_ident.add_row().cells
        style_cell(r[0], label, bold=True); set_cell_background(r[0], "D1F2EB")
        items = safe_list(ident_src.get(key)); r[1].text = ""
        for idx, item in enumerate(items):
            p = r[1].paragraphs[0] if idx == 0 else r[1].add_paragraph()
            p.add_run(item)
            if key not in ["dimensi_profil", "panca_cinta", "latar_belakang"]: p.style = 'List Bullet'
    doc.add_paragraph()

    doc.add_paragraph("B. DESAIN PEMBELAJARAN").bold = True
    t_desain = doc.add_table(rows=0, cols=2); t_desain.style = 'Table Grid'
    t_desain.columns[0].width = Cm(5.0); t_desain.columns[1].width = Cm(13.0)
    desain_src = d1.get("desain", {}) if isinstance(d1.get("desain"), dict) else {}
    for label, key in [("Capaian Pembelajaran (CP)", "capaian_pembelajaran"), ("Tujuan Pembelajaran (TP)", "tujuan_pembelajaran"), ("Lintas Disiplin Ilmu", "lintas_disiplin"), ("Topik Pembelajaran", "topik_pembelajaran"), ("Praktik Pedagogi", "praktik_pedagogi"), ("Lingkungan Belajar", "lingkungan_belajar"), ("Kemitraan Pembelajaran", "kemitraan_pembelajaran"), ("Pemanfaatan Digital", "pemanfaatan_digital")]:
        r = t_desain.add_row().cells
        style_cell(r[0], label, bold=True); set_cell_background(r[0], "FADBD8")
        items = safe_list(desain_src.get(key)); r[1].text = ""
        for idx, item in enumerate(items):
            p = r[1].paragraphs[0] if idx == 0 else r[1].add_paragraph()
            p.add_run(item)
            if key not in ["capaian_pembelajaran", "tujuan_pembelajaran"]: p.style = 'List Bullet'
    doc.add_paragraph()

    doc.add_heading("PENGALAMAN BELAJAR", level=1)
    pertemuan_list = safe_list(d2.get("pertemuan", []))
    for p in pertemuan_list:
        if not isinstance(p, dict): continue
        no_pert = p.get("nomor", "1"); materi = p.get("materi", "Materi"); durasi = p.get("durasi", "1 x 35 Menit")
        doc.add_paragraph(f"PENGALAMAN BELAJAR – PERTEMUAN {no_pert}").bold = True
        doc.add_paragraph(f"Materi: {materi}\nDurasi: {durasi} | Model: PBL | Metode: Ceramah Interaktif, Diskusi Kelompok, Tanya Jawab")
        
        t_pem = doc.add_table(rows=2, cols=4); t_pem.style = 'Table Grid'
        t_pem.columns[0].width = Cm(4.0); t_pem.columns[1].width = Cm(9.0); t_pem.columns[2].width = Cm(2.0); t_pem.columns[3].width = Cm(3.0)
        h_pem = t_pem.rows[0].cells
        for i, txt in enumerate(["FASE KEGIATAN", "AKTIVITAS PEMBELAJARAN", "WAKTU", "PRINSIP DL"]): style_cell(h_pem[i], txt, bold=True, center=True); set_cell_background(h_pem[i], "FCF3CF")
        r_pem = t_pem.rows[1].cells
        r_pem[0].text = "PEMBUKAAN\n(Meaningful – Bermakna)"
        pem_data = p.get("pembukaan", {}) if isinstance(p.get("pembukaan"), dict) else {}
        akt_pem = "\n".join([f"{i+1}. {a}" for i, a in enumerate(safe_list(pem_data.get("aktivitas")))])
        r_pem[1].text = akt_pem
        style_cell(r_pem[2], pem_data.get("waktu", "10'"), center=True); style_cell(r_pem[3], "MEANINGFUL\n(Bermakna)", center=True)
        doc.add_paragraph()

        t_inti = doc.add_table(rows=4, cols=4); t_inti.style = 'Table Grid'
        t_inti.columns[0].width = Cm(4.0); t_inti.columns[1].width = Cm(4.5); t_inti.columns[2].width = Cm(6.5); t_inti.columns[3].width = Cm(3.0)
        h_inti = t_inti.rows[0].cells
        for i, txt in enumerate(["PENGALAMAN BELAJAR", "SINTAK PBL", "AKTIVITAS PEMBELAJARAN (KEGIATAN INTI)", "PRINSIP DL"]): style_cell(h_inti[i], txt, bold=True, center=True); set_cell_background(h_inti[i], "E8DAEF")
        inti_phases = [("MEMAHAMI", p.get("inti_memahami", {}), "MINDFUL\n(Berkesadaran)"), ("MENGAPLIKASIKAN", p.get("inti_mengaplikasikan", {}), "JOYFUL\n(Menggembirakan)"), ("MEREFLEKSIKAN", p.get("inti_merefleksikan", {}), "MEANINGFUL\n(Bermakna)")]
        for idx, (nama_fase, data_fase, prinsip) in enumerate(inti_phases, 1):
            if not isinstance(data_fase, dict): data_fase = {}
            r_i = t_inti.rows[idx].cells
            style_cell(r_i[0], nama_fase, bold=True); r_i[1].text = data_fase.get("sintak_pbl", "Langkah PBL")
            akt_i = "\n".join([f"• {a}" for a in safe_list(data_fase.get("aktivitas"))]); r_i[2].text = akt_i; style_cell(r_i[3], prinsip, center=True)
        doc.add_paragraph()
        
        t_pen = doc.add_table(rows=2, cols=4); t_pen.style = 'Table Grid'
        t_pen.columns[0].width = Cm(4.0); t_pen.columns[1].width = Cm(9.0); t_pen.columns[2].width = Cm(2.0); t_pen.columns[3].width = Cm(3.0)
        h_pen = t_pen.rows[0].cells
        for i, txt in enumerate(["FASE KEGIATAN", "AKTIVITAS PEMBELAJARAN", "WAKTU", "PRINSIP DL"]): style_cell(h_pen[i], txt, bold=True, center=True); set_cell_background(h_pen[i], "D1F2EB")
        r_pen = t_pen.rows[1].cells
        r_pen[0].text = "PENUTUP\n(Mindful – Berkesadaran)"
        pen_data = p.get("penutup", {}) if isinstance(p.get("penutup"), dict) else {}
        akt_pen = "\n".join([f"{i+1}. {a}" for i, a in enumerate(safe_list(pen_data.get("aktivitas")))])
        r_pen[1].text = akt_pen
        style_cell(r_pen[2], pen_data.get("waktu", "10'"), center=True); style_cell(r_pen[3], "MINDFUL\n(Berkesadaran)", center=True)
        doc.add_paragraph()

    doc.add_paragraph("PENILAIAN / ASESMEN").bold = True
    t_nilai = doc.add_table(rows=0, cols=2); t_nilai.style = 'Table Grid'
    t_nilai.columns[0].width = Cm(5.0); t_nilai.columns[1].width = Cm(13.0)
    pen_src = d3.get("penilaian", {}) if isinstance(d3.get("penilaian"), dict) else {}
    
    r_n1 = t_nilai.add_row().cells
    style_cell(r_n1[0], "Asesmen Awal (Diagnostik)", bold=True); set_cell_background(r_n1[0], "D6EAF8")
    r_n1[1].text = "\n".join([f"• {a}" for a in safe_list(pen_src.get("awal"))])
    r_n2 = t_nilai.add_row().cells
    style_cell(r_n2[0], "Asesmen Formatif (As Learning & For Learning)", bold=True); set_cell_background(r_n2[0], "D6EAF8")
    as_learn = "\n".join([f"• {a}" for a in safe_list(pen_src.get("formatif_as"))])
    for_learn = "\n".join([f"• {a}" for a in safe_list(pen_src.get("formatif_for"))])
    r_n2[1].text = f"As Learning:\n{as_learn}\n\nFor Learning:\n{for_learn}"
    r_n3 = t_nilai.add_row().cells
    style_cell(r_n3[0], "Asesmen Sumatif (Of Learning)", bold=True); set_cell_background(r_n3[0], "D6EAF8")
    r_n3[1].text = "\n".join([f"• {a}" for a in safe_list(pen_src.get("sumatif"))])
    doc.add_paragraph()

    doc.add_paragraph("LAMPIRAN I – ASESMEN").bold = True
    doc.add_paragraph("A. ASESMEN AWAL (LISAN)").bold = True
    for i, a in enumerate(safe_list(d3.get("asesmen_awal_lisan")), 1): doc.add_paragraph(f"{i}. {a}")
    doc.add_paragraph()
    
    doc.add_paragraph("B. RUBRIK PENILAIAN SIKAP (Skala 1–4)").bold = True
    t_sikap = doc.add_table(rows=5, cols=5); t_sikap.style = 'Table Grid'
    h_sikap = ["Aspek Sikap", "Skor 4 (Sangat Baik)", "Skor 3 (Baik)", "Skor 2 (Cukup)", "Skor 1 (Perlu Bimb.)"]
    for i, h in enumerate(h_sikap): style_cell(t_sikap.cell(0, i), h, bold=True, center=True); set_cell_background(t_sikap.cell(0, i), "FCF3CF")
    sikap_data = [["Disiplin", "Selalu hadir tepat waktu dan mengikuti semua aturan", "Hadir tepat waktu, sedikit pelanggaran kecil", "Sering terlambat atau melanggar beberapa aturan", "Sering tidak hadir tanpa keterangan"], ["Tanggung Jawab", "Menyelesaikan tugas tepat waktu dan berkualitas", "Menyelesaikan tugas meski ada kekurangan kecil", "Tugas sering tidak selesai atau terlambat", "Tugas tidak dikerjakan"], ["Kerjasama", "Aktif berkontribusi dan mendorong anggota lain", "Berkontribusi dalam kelompok secara aktif", "Sesekali berkontribusi jika diminta", "Tidak berkontribusi dalam kelompok"], ["Toleransi", "Sangat menghargai pendapat semua teman", "Menghargai pendapat dengan baik", "Sesekali kurang menghargai pendapat", "Tidak menghargai pendapat teman"]]
    for r_idx, row_data in enumerate(sikap_data, start=1):
        for c_idx, cell_data in enumerate(row_data): style_cell(t_sikap.cell(r_idx, c_idx), cell_data)
    doc.add_paragraph()
    
    doc.add_paragraph("C. RUBRIK PENILAIAN PENGETAHUAN (Skala 1–4)").bold = True
    t_penget = doc.add_table(rows=1, cols=5); t_penget.style = 'Table Grid'
    h_penget = ["Aspek", "Skor 4", "Skor 3", "Skor 2", "Skor 1"]
    for i, h in enumerate(h_penget): style_cell(t_penget.cell(0, i), h, bold=True, center=True); set_cell_background(t_penget.cell(0, i), "FCF3CF")
    for rubrik in safe_list(d3.get("rubrik_pengetahuan", [])):
        if not isinstance(rubrik, dict): continue
        r = t_penget.add_row().cells
        style_cell(r[0], rubrik.get("aspek", "Aspek"), bold=True)
        r[1].text = rubrik.get("skor_4", ""); r[2].text = rubrik.get("skor_3", ""); r[3].text = rubrik.get("skor_2", ""); r[4].text = rubrik.get("skor_1", "")
    doc.add_paragraph()
    
    doc.add_paragraph("D. ASESMEN SUMATIF – 5 SOAL URAIAN HOTS").bold = True
    for i, a in enumerate(safe_list(d3.get("sumatif_hots")), 1): doc.add_paragraph(f"{i}. {a}")
    doc.add_paragraph()

    doc.add_paragraph("LAMPIRAN II – MATERI AJAR").bold = True
    for mat in safe_list(d3.get("materi_ajar", [])):
        if not isinstance(mat, dict): continue
        doc.add_paragraph(mat.get("sesi", "SESI")).bold = True
        doc.add_paragraph(f"Tujuan: {mat.get('tujuan', '')}")
        t_mat = doc.add_table(rows=0, cols=2); t_mat.style = 'Table Grid'
        t_mat.columns[0].width = Cm(4.0); t_mat.columns[1].width = Cm(14.0)
        for k, v in [("Pengantar", mat.get("pengantar")), ("Konsep Utama", mat.get("konsep_utama")), ("Contoh Kontekstual", mat.get("contoh"))]:
            r = t_mat.add_row().cells
            style_cell(r[0], k, bold=True); set_cell_background(r[0], "FADBD8")
            r[1].text = str(v)
        doc.add_paragraph()

    doc.add_paragraph("LAMPIRAN V – TINDAK LANJUT DAN REFLEKSI").bold = True
    doc.add_paragraph("A. PROGRAM REMEDIAL").bold = True
    doc.add_paragraph(f"Mata Pelajaran: {form['mapel']} | Kelas: {form['kelas']} | Semester: {form['semester']}")
    doc.add_paragraph("Bentuk Asesmen: Tes Tulis\nBentuk Pelaksanaan Remedial:\n• Bimbingan Individu: untuk siswa dengan nilai < 60\n• Latihan Ulang Terbimbing: untuk siswa dengan nilai 60–74")
    doc.add_paragraph("5 Soal Pilihan Ganda Remedial:").bold = True
    for i, pg in enumerate(safe_list(d3.get("remedial_pg", [])), 1):
        if not isinstance(pg, dict): continue
        doc.add_paragraph(f"{i}. {pg.get('soal', '')}")
        doc.add_paragraph(f"   a. {pg.get('a', '')}\n   b. {pg.get('b', '')}\n   c. {pg.get('c', '')}\n   d. {pg.get('d', '')}")
        doc.add_paragraph(f"   → Jawaban: {pg.get('kunci', '')}")
    doc.add_paragraph()
    
    doc.add_paragraph("B. PROGRAM PENGAYAAN").bold = True
    doc.add_paragraph("Bentuk Pengayaan: Proyek Mini & Esai HOTS\nUntuk peserta didik yang telah menguasai semua TP:")
    for i, p in enumerate(safe_list(d3.get("pengayaan", [])), 1): doc.add_paragraph(f"{i}. {p}")
    doc.add_paragraph()
    
    doc.add_paragraph("C. INSTRUMEN REFLEKSI PEMBELAJARAN").bold = True
    t_ref = doc.add_table(rows=1, cols=2); t_ref.style = 'Table Grid'
    t_ref.columns[0].width = Cm(9.0); t_ref.columns[1].width = Cm(9.0)
    style_cell(t_ref.cell(0,0), "Refleksi Peserta Didik", bold=True); set_cell_background(t_ref.cell(0,0), "D1F2EB")
    style_cell(t_ref.cell(0,1), "Refleksi Guru", bold=True); set_cell_background(t_ref.cell(0,1), "D1F2EB")
    r_ref2 = t_ref.add_row().cells
    r_ref2[0].text = "1. Apa yang sudah aku pahami dari pembelajaran ini?\n   ________________________________\n2. Apa yang masih terasa sulit bagiku?\n   ________________________________\n3. Bagaimana perasaanku selama belajar hari ini? (Sangat Baik / Cukup / Perlu Perbaikan)\n4. Bagaimana aku bisa menerapkan yang dipelajari?\n   ________________________________"
    r_ref2[1].text = "1. Apakah tujuan pembelajaran tercapai hari ini?\n   ________________________________\n2. Strategi apa yang paling efektif/kurang efektif?\n   ________________________________\n3. Kendala yang muncul selama pembelajaran:\n   ________________________________\n4. Rencana perbaikan untuk pertemuan berikutnya:\n   ________________________________"
    doc.add_paragraph()

    doc.add_paragraph("GLOSARIUM").bold = True
    t_glo = doc.add_table(rows=1, cols=2); t_glo.style = 'Table Grid'
    t_glo.columns[0].width = Cm(4.0); t_glo.columns[1].width = Cm(14.0)
    style_cell(t_glo.cell(0,0), "ISTILAH", bold=True, center=True); set_cell_background(t_glo.cell(0,0), "E8DAEF")
    style_cell(t_glo.cell(0,1), "DEFINISI", bold=True, center=True); set_cell_background(t_glo.cell(0,1), "E8DAEF")
    for glo in safe_list(d3.get("glosarium", [])):
        if not isinstance(glo, dict): continue
        r = t_glo.add_row().cells
        style_cell(r[0], glo.get("istilah", ""), bold=True)
        r[1].text = glo.get("definisi", "")
    doc.add_paragraph()

    doc.add_paragraph("DAFTAR PUSTAKA").bold = True
    for i, dp in enumerate(safe_list(d3.get("daftar_pustaka", [])), 1): doc.add_paragraph(f"{i}. {dp}")

    add_signatures(doc, form, full_width=True)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


def build_lkpd_per_bab(form, bab_name, d4_lkpd_data):
    doc = create_base_doc(landscape=False)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LEMBAR KERJA PESERTA DIDIK (LKPD)\n")
    r.bold = True
    r.font.size = Pt(14)
    p.add_run(f"{form['mapel']} Kelas {form['kelas']} - {bab_name}")
    
    doc.add_paragraph("Nama\t\t: __________________________")
    doc.add_paragraph("Kelas\t\t: __________________________")
    doc.add_paragraph("Tanggal\t: __________________________")
    doc.add_paragraph()
    
    lkpd_list = safe_list(d4_lkpd_data.get("lkpd", []))
    for item in lkpd_list:
        if not isinstance(item, dict): continue
        topik_judul = str(item.get("topik_judul", f"Topik: {item.get('topik', 'Aktivitas Belajar')}"))
        doc.add_heading(topik_judul, level=2)
        if item.get("tujuan_pembelajaran"):
            doc.add_heading("Tujuan Pembelajaran:", level=3)
            for tp in safe_list(item.get("tujuan_pembelajaran")): doc.add_paragraph(tp, style='List Bullet')
            doc.add_paragraph()
            
        aktivitas_list = safe_list(item.get("aktivitas", []))
        for act in aktivitas_list:
            if not isinstance(act, dict): continue
            judul_act = act.get("judul_aktivitas", "Aktivitas")
            doc.add_heading(judul_act, level=3)
            if act.get("konteks"): doc.add_paragraph(str(act.get("konteks")))
            if act.get("ayo_berpikir"):
                doc.add_heading("Ayo Menggambar dan Berpikir!", level=4)
                for ab in safe_list(act.get("ayo_berpikir")): doc.add_paragraph(ab, style='List Bullet')
            if act.get("koneksi_matematika"):
                doc.add_heading("Menghubungkan ke Konsep/Materi:", level=4)
                for kc in safe_list(act.get("koneksi_matematika")): doc.add_paragraph(kc, style='List Bullet')
            if act.get("kesimpulan"):
                p_kes = doc.add_paragraph(); p_kes.add_run("Kesimpulanmu: ").bold = True; p_kes.add_run(str(act.get("kesimpulan"))); doc.add_paragraph("\n") 
            if act.get("latihan_berpikir"):
                p_lat = doc.add_paragraph(); p_lat.add_run("Latihan Berpikir:\n").bold = True; p_lat.add_run(str(act.get("latihan_berpikir"))); doc.add_paragraph("\n\n") 
        doc.add_page_break()
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()

# ==============================================================================
# UI STREAMLIT 
# ==============================================================================
st.title("📘 MI MIFTAHUSSALAM ADMIN GURU GENERATOR V.4.13 ")
st.markdown("*Berbasis Model Lagos AI 9.1 - Joyful Colors & Full Dokumen*")

with st.form("form_modul"):
    col1, col2 = st.columns(2)
    with col1:
        mapel = st.text_input("Mata Pelajaran", placeholder="Contoh: Matematika")
        kelas = st.selectbox("Jenjang / Kelas", list(JENJANG_FASE.keys()), index=6) 
        semester = st.selectbox("Semester", ["1 (Ganjil)", "2 (Genap)", "1 & 2 (Satu Tahun Penuh)"])
    with col2:
        alokasi = st.text_input("Alokasi Waktu per Pertemuan", value="4 JP x 35 menit")
        sekolah = st.text_input("Sekolah", value="MI Miftahussalam")
        tahun_pelajaran = st.text_input("Tahun Pelajaran", value="2026/2027")

    st.divider()
    st.markdown("### 📚 Data Bab Pembelajaran & Sub-Bab")
    st.info("Ketik nama Bab beserta rincian Sub-Bab di dalam kurung, lalu jumlah pertemuannya. Pisahkan dengan tanda sama dengan (=).")
    
    default_bab_text = """Bab 1: Bangga Menjadi Anak Indonesia (Sub-bab: A. Aku Anak Indonesia, B. Pancasila di Hatiku) = 2
Bab 2: Musisi Indonesia di Pentas Dunia (Sub-bab: A. Mengenal Alat Musik, B. Wawancara Tokoh) = 2"""
    
    input_bab_raw = st.text_area("Daftar Bab & Jumlah Pertemuan", value=default_bab_text, height=120)

    st.divider()
    col3, col4 = st.columns(2)
    now = datetime.datetime.now()
    bulan_indo = ["Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
    titimangsa_otomatis = f"Bogor, {now.day} {bulan_indo[now.month - 1]} {now.year}"

    with col3:
        titimangsa = st.text_input("Titimangsa", value=titimangsa_otomatis)
        penyusun = st.text_input("Penyusun (Guru)", placeholder="Erian Kurniawan, S.E.")
    with col4:
        kepala_madrasah = st.text_input("Kepala Madrasah", placeholder="Drs. Andi Supriadi")

    st.divider()
    
    jenis_cover = st.selectbox("Pilih Jenis Cover", ["Tanpa Cover", "Cover Administrasi Guru (Buku Umum)", "Cover Modul Ajar", "Cover Program Tahunan & Semester", "Cover Jurnal Mengajar", "Cover CP", "Cover ATP", "Cover KKTP"])
    pilihan_dokumen = st.multiselect(
        "Pilih dokumen yang ingin di-generate otomatis",
        ["Modul Ajar (Per Bab)", "LKPD Siswa (Per Bab)", "Capaian Pembelajaran (CP)", "Alur Tujuan Pembelajaran (ATP)", "Prota", "Promes", "KKTP", "Jurnal Mengajar (Global)"],
        default=["Modul Ajar (Per Bab)", "LKPD Siswa (Per Bab)", "Capaian Pembelajaran (CP)", "Promes", "KKTP"]
    )

    submitted = st.form_submit_button("✨ Eksekusi & Generate Semua Dokumen", use_container_width=True)

if submitted:
    daftar_bab = []
    if input_bab_raw.strip():
        for line in input_bab_raw.strip().split('\n'):
            if "=" in line:
                parts = line.split("=")
                try: jml_pert = int(parts[1].strip())
                except ValueError: jml_pert = 2 
                daftar_bab.append({"Bab / Topik": parts[0].strip(), "Jumlah Pertemuan": jml_pert})
            elif line.strip(): daftar_bab.append({"Bab / Topik": line.strip(), "Jumlah Pertemuan": 2})

    if not (mapel and penyusun and sekolah and kepala_madrasah): st.warning("Lengkapi data identitas (Mapel, Penyusun, Sekolah, Kepala).")
    elif len(daftar_bab) == 0: st.warning("Harap isi setidaknya 1 Bab pada kotak input teks.")
    else:
        form = dict(mapel=mapel, kelas=kelas, semester=semester, alokasi=alokasi, penyusun=penyusun, sekolah=sekolah, tahun_pelajaran=tahun_pelajaran, titimangsa=titimangsa, kepala_madrasah=kepala_madrasah)
        safe_mapel = re.sub(r'[^a-zA-Z0-9_\-]', '_', form['mapel']); safe_kelas = re.sub(r'[^a-zA-Z0-9_\-]', '_', form['kelas'].split()[0])
        st.session_state["hasil_generate"] = {}
        
        try:
            if jenis_cover != "Tanpa Cover":
                cover_status = st.empty(); cover_status.info(f"📄 Membuat {jenis_cover}...")
                st.session_state["hasil_generate"][f"{jenis_cover}_{safe_mapel}.docx"] = build_cover(form, jenis_cover)
                cover_status.success(f"✅ Selesai membuat {jenis_cover}")

            all_chapters_data = []; total_pertemuan_global = 0
            
            st.markdown("### 🧠 Proses Pembuatan Master Data, Modul & LKPD")
            for idx, row in enumerate(daftar_bab):
                bab_name = str(row["Bab / Topik"]); jml_pert = row["Jumlah Pertemuan"]
                total_pertemuan_global += jml_pert; safe_bab = re.sub(r'[^a-zA-Z0-9_\-]', '_', bab_name)
                
                st.write(f"🔄 **Sedang memproses {bab_name} ({jml_pert} Pertemuan)...**")
                
                st.info("Tahap 1: Desain Pembelajaran Modul..."); d1_ctx = call_ai(prompt_step_1(form, bab_name)); time.sleep(3)
                st.info("Tahap 2: Pengalaman Belajar Modul..."); d2_ctx = call_ai(prompt_step_2(form, bab_name, jml_pert, d1_ctx)); time.sleep(3)
                st.info("Tahap 3: Asesmen, Rubrik & Lampiran Modul..."); d3_ctx = call_ai(prompt_step_3(form, bab_name, jml_pert, d2_ctx))
                
                chapter_data = {"bab": bab_name, "jml_pert": jml_pert, "d1": d1_ctx, "d2": d2_ctx, "d3": d3_ctx}
                all_chapters_data.append(chapter_data)
                
                if "Modul Ajar (Per Bab)" in pilihan_dokumen:
                    st.session_state["hasil_generate"][f"Modul_Ajar_{safe_mapel}_Bab_{idx+1}_{safe_bab}.docx"] = build_modul_ajar(form, bab_name, jml_pert, chapter_data)
                    st.success(f"✅ Berhasil membuat Modul Ajar: {bab_name}")
                    
                if "LKPD Siswa (Per Bab)" in pilihan_dokumen:
                    combined_lkpd_list = []
                    for p_ke in range(1, jml_pert + 1):
                        st.info(f"Tahap Khusus: Pembuatan LKPD Interaktif (Pertemuan {p_ke}/{jml_pert})...")
                        lkpd_resp = call_ai(prompt_lkpd_interaktif_per_pertemuan(form, bab_name, p_ke))
                        raw_lkpd = lkpd_resp.get("lkpd")
                        if isinstance(raw_lkpd, list) and len(raw_lkpd) > 0: combined_lkpd_list.append(raw_lkpd[0])
                        elif isinstance(raw_lkpd, dict): combined_lkpd_list.append(raw_lkpd)
                        time.sleep(3)
                        
                    d4_combined_ctx = {"lkpd": combined_lkpd_list}
                    st.session_state["hasil_generate"][f"LKPD_Siswa_{safe_mapel}_Bab_{idx+1}_{safe_bab}.docx"] = build_lkpd_per_bab(form, bab_name, d4_combined_ctx)
                    st.success(f"✅ Berhasil membuat LKPD Interaktif: {bab_name}")
                
                time.sleep(5) 
            
            master_context = get_aggregated_sinkronisasi_context(all_chapters_data)
            st.markdown("### ⚙️ Pemrosesan Dokumen Administrasi Global (Sinkronisasi)")
            
            for tipe in [d for d in pilihan_dokumen if d not in ["Modul Ajar (Per Bab)", "LKPD Siswa (Per Bab)"]]:
                doc_status = st.empty()
                try:
                    if tipe == "Jurnal Mengajar (Global)":
                        doc_status.info("🔄 Menyusun Jurnal Mengajar Harian (Otomatis dari Modul)...")
                        jurnal_data = generate_jurnal_otomatis(form, all_chapters_data)
                        doc_bytes = build_jurnal(form, jurnal_data)
                        
                    elif tipe == "KKTP":
                        doc_status.info("🔄 Menyusun KKTP secara bertahap (Per Bab) agar tidak terpotong...")
                        kktp_combined_rows = []
                        for chap in all_chapters_data:
                            bab_nm = chap["bab"]
                            tp_list = safe_list(chap.get("d1", {}).get("desain", {}).get("tujuan_pembelajaran", []))
                            st.info(f"Generate KKTP: {bab_nm}...")
                            kktp_resp = call_ai(prompt_kktp_per_bab(form, bab_nm, tp_list))
                            kktp_combined_rows.extend(safe_list(kktp_resp.get("rows", [])))
                            time.sleep(3)
                        doc_bytes = build_kktp(form, {"rows": kktp_combined_rows})
                        
                    else:
                        doc_status.info(f"🔄 Mensinkronkan & memproses: {tipe} (Sesuai KMA 1503)...")
                        if tipe == "Capaian Pembelajaran (CP)": doc_bytes = build_cp(form, call_ai(prompt_cp(form, master_context)))
                        elif tipe == "Alur Tujuan Pembelajaran (ATP)": doc_bytes = build_atp(form, call_ai(prompt_atp(form, master_context)))
                        elif tipe == "Prota": doc_bytes = build_prota(form, call_ai(prompt_prota(form, master_context)))
                        elif tipe == "Promes": doc_bytes = build_promes(form, call_ai(prompt_promes(form, master_context)))
                    
                    st.session_state["hasil_generate"][f"{tipe.replace(' & ', '_').replace(' ', '_').replace('(', '').replace(')', '')}_{safe_mapel}_Sinkronisasi.docx"] = doc_bytes
                    doc_status.success(f"✅ Berhasil membuat & mensinkronkan: {tipe}")
                    time.sleep(8) 
                except Exception as e: doc_status.error(f"❌ Gagal memproses {tipe}: {e}")
            
            st.success("🎉 Selesai! Seluruh Ekosistem Dokumen Administrasi Anda Berhasil Dibuat.")
        except Exception as e: st.error(f"Terjadi kesalahan sistem: {e}")

if "hasil_generate" in st.session_state and st.session_state["hasil_generate"]:
    st.divider()
    st.write("⬇️ **Silakan Unduh File Anda di Bawah Ini:**")
    for fname, fbytes in st.session_state["hasil_generate"].items():
        st.download_button(label=f"Unduh {fname}", data=fbytes, file_name=fname, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
